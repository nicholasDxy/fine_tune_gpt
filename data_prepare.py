import docx
import json
import pandas as pd
import jsonlines

# Read the docx file


# Extract the QA pairs
questions_title = []
questions_description = []
answers = []
template_user_content = ''
template_assistant_content = ''
template = {"messages": [{"role": "system", "content": " As a helpful assistant, your task is to write an impressive personal statement for a student who is applying for a study program. You will be provided with a title, requirements, and the user's experience to base your essay on. Your goal is to showcase the student's good character in a polished and compelling manner. To accurately reflect the author's point of view, you should use a first-person perspective and incorporate the word \"I\" throughout the essay."}, 
                         {"role": "user", "content": ''}, 
                         {"role": "assistant", "content": ''}]}

def parse_doc():
    doc = docx.Document('data/data.docx')
    temp_answer = None
    temp_title = None
    temp_description = None
    for i in range(len(doc.paragraphs)):
        for j in range(len(doc.paragraphs[i].runs)):
            run = doc.paragraphs[i].runs[j]
            if run.font.size == 381000:
                if temp_answer != None:
                    answers.append(temp_answer)
                    temp_answer = None
                if temp_title == None:
                    temp_title = run.text
                else:
                    temp_title += run.text
            elif run.font.size == 203200:
                if temp_title != None:
                    questions_title.append(temp_title)
                    temp_title = None
                if temp_description == None:
                    temp_description = run.text
                else:
                    temp_description += run.text
            elif run.font.size == 152400:
                if temp_description != None:
                    questions_description.append(temp_description)
                    temp_description = None
                if temp_answer == None:
                    temp_answer = run.text
                else:
                    temp_answer += run.text

    if temp_answer !=None:
        answers.append(temp_answer)
        temp_answer = None
    
    print(len(questions_title))
    # print(questions_title)
    print(len(questions_description))
    print(len(answers))
    
    with jsonlines.open('data/doc_info.jsonl',mode='a') as writer:
        for i in range(0,10):
            print(questions_description[i])
            info = {}
            info['title'] = questions_title[i]
            info['description'] = questions_description[i]
            info['content'] = answers[i]
            writer.write(info)


def generate_train_data():
    doc_list = []
    message_list = []
    with open("data/summary_info_new.jsonl", "r+", encoding="utf8") as f:
        for item in jsonlines.Reader(f):
            doc_list.append(item)
            
    with jsonlines.open('data/train_data.jsonl',mode='a') as writer:
        for item in doc_list:
            template_user_content = "Title: " + item['title'] + ". Requirement : " + item['description'] + '. Experience: ' + item['summary']
            template["messages"][1]["content"] = template_user_content
            template["messages"][2]["content"] = item['content']
            writer.write(template)
        
# Write the QA pairs to a csv file
# df = pd.DataFrame({'Question_title': questions_title,'questions_description':questions_description, 'Answer': answers})
# df.to_csv('train_data.csv', index=False)

# with jsonlines.open('data/train_data.jsonl',mode='a') as writer:
#     for i in range(0,10):
#         print(questions_description[i])
#         template_user_content = "Title: " + questions_title[i] + ". Description: " + questions_description[i]
#         template_assistant_content = answers[i]
#         template["messages"][1]["content"] = template_user_content
#         template["messages"][2]["content"] = template_assistant_content
#         writer.write(template)
    
generate_train_data()
